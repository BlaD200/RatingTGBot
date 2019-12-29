import os
import re

from docx import Document
from openpyxl import load_workbook

from sql.subjects_info_writer import fill_subjects_info_table, get_faculty, is_mp
from timetable import timetables_dir, timetable_downloader


def extract_from_xlsx(filename: str):
    """
    Parser for <filename>.xlsx timetable
    Process the timetable and extract speciality, studding year and subjects.
    Also access to :code:`update_info` table to get faculty and to know if this is timetable for master's program

    :param filename: timetable filename to parse
    """
    wb = load_workbook(filename="%s/%s" % (timetables_dir, filename))
    ws = wb.copy_worksheet(wb.active)

    info = ws['A7'].value
    faculty = ws['A6'].value
    speciality = info[info.find('"') + 1:info.find('"', info.find('"') + 1)]
    year = [int(ch) for ch in info if ch.isdigit()][0]

    subjects = list()
    for row in ws.iter_rows(11, 72, max_col=7):
        subject = row[2].value
        if subject:
            subject = subject.split(',')[0]
            if subject not in subjects:
                subjects.append(subject)
    fill_subjects_info_table(subjects, year, speciality, faculty)


def extract_from_docx(filename: str):
    """
    Parser for <filename>.docx timetable
    Process the timetable and extract speciality, studding year and subjects.
    Also refer to :code:`update_info` table to get faculty and to know if this is timetable for master's program

    :param filename: timetable filename to parse
    """
    pattern = r'([с|С|c]?.*ність)([\sа-яА-ЯіїІ"«»(),\-:;_]*?)([\s,МП]*)([\s_\-]*)([1-4]+?)([\sрн\.]+)'
    doc = Document("%s/%s" % (timetables_dir, filename))

    for table in doc.tables:
        faculty_name = get_faculty(filename)
        info = ''.join([paragraph.text for paragraph in doc.paragraphs])
        m = re.search(pattern, info)
        speciality = m.group(2)
        year = int(m.group(5))
        if is_mp(filename):
            year += 4
        for symb in ['"', "'", '«', '»', '_']:
            speciality = speciality.replace(symb, '').strip()

        subjects = list()
        for row in table.rows[1:]:
            if len(row.cells) > 2:
                subject = [run.text for run in row.cells[2].paragraphs[0].runs if run.bold and run.text]
                if subject:
                    subject = ''.join(subject)
                    subject = subject.strip()
                    if subject and subject not in subjects:
                        subjects.append(subject)
        fill_subjects_info_table(subjects, year, speciality, faculty_name)


def process_file(filename: str):
    """
    Manage to call correct parser function for giving file with filename

    :param filename: file name the will be processed
    """
    if filename.endswith('.xlsx'):
        extract_from_xlsx(filename)
    elif filename.endswith('.docx'):
        extract_from_docx(filename)
    elif filename.endswith('.doc'):
        print("Removing...    ||{:<100s}||".format(filename))
        os.remove('%s/%s' % (timetables_dir, filename))
        return

    print("Parsed...     ||{:<100s}||".format(filename))
    os.remove('%s/%s' % (timetables_dir, filename))


def extract_from_files():
    """
    Manage to process each file in /timetables directory or has just downloaded via
    :code:`timetables.timetable_downloader` function
    """
    # TODO some subjects didn't mark bold (Екологія БП-1)

    if os.listdir(timetables_dir):
        for filename in os.listdir(timetables_dir):
            process_file(filename)

    for filename in timetable_downloader.update_timetables():
        process_file(filename)


if __name__ == '__main__':
    extract_from_files()
