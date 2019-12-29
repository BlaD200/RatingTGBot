from docx import Document
from openpyxl import load_workbook

import os
import re

from timetable import timetable_downloader
from timetable import timetables_dir

from sql import initialize_connection, update_last_index
from mysql.connector.errors import DataError


def fill_subjects_info_table(subjects: list, year: int, speciality: str, faculty: str):
    """
    Accesses to :code:`subjects_info` in the database to populate table with given values

    :param subjects: list of str of contains subjects names
    :param year: year of study
    :param speciality: name of speciality
    :param faculty: name of faculty
    """
    db = initialize_connection()
    cursor = db.cursor()

    update_last_index(db, "subjects_info", "subject_id")

    for subject in subjects:
        values = (subject, year, speciality, faculty)
        # print(values, sep=',')
        for val in values:
            if not val:
                var_name = [k for k, v in locals().items() if v == val][0]
                raise ValueError('%s cannot be "%s"' % (var_name, val))
        cursor.execute(
            "SELECT subject_id from subjects_info WHERE subject_name = %s AND "
            "study_year = %s AND faculty_name = %s;", (subject, year, faculty))
        old_id = cursor.fetchall()
        try:
            if old_id:
                old_id = old_id[0][0]
                cursor.execute("INSERT INTO subjects_info VALUES"
                               "(%s, %s, %s, %s, %s) ON DUPLICATE KEY UPDATE "
                               "subject_name = %s, study_year = %s, specialty = %s, faculty_name = %s",
                               (old_id, *values, *values))
            else:
                cursor.execute("INSERT INTO subjects_info(subject_name, study_year, specialty, faculty_name) VALUES"
                               "(%s, %s, %s, %s)", values)
        except DataError as e:
            print(values)
            raise DataError

        db.commit()

    db.close()


def get_faculty(timetable_name):
    """
    Access to :code:`update_info` table to get the faculty name

    :param timetable_name: filename of timetable
    :return: faculty name
    """
    db = initialize_connection()
    cursor = db.cursor()

    cursor.execute("SELECT faculty_name from update_info WHERE timetable_name = %s;", (timetable_name,))
    faculty_name = cursor.fetchone()
    db.close()

    return faculty_name[0]


def is_mp(timetable_name: str) -> bool:
    """
    Access to :code:`update_info` table to get if this is timetable for master's program

    :param timetable_name: filename of timetable
    :return: True if given timetable is for master's program else False
    """
    db = initialize_connection()
    cursor = db.cursor()

    cursor.execute("SELECT isMP FROM update_info WHERE timetable_name = %s;", (timetable_name,))
    mp = cursor.fetchone()
    db.close()

    return mp[0]


def extract_from_xlsx(filename: str):
    """
    Parser for <filename>.xlsx timetable
    Process the timetable and extract speciality, studding year and subjects.
    Also access to :code:`update_info` table to get faculty and to know if this is timetable for master's program

    :param filename: timetable filename to parse
    """
    wb = load_workbook(filename="%s/%s" % (timetables_dir, filename))
    ws = wb.copy_worksheet(wb.active)

    info = ws['A7'].value
    faculty = ws['A6'].value
    speciality = info[info.find('"') + 1:info.find('"', info.find('"') + 1)]
    year = [int(ch) for ch in info if ch.isdigit()][0]

    subjects = list()
    for row in ws.iter_rows(11, 72, max_col=7):
        subject = row[2].value
        if subject:
            subject = subject.split(',')[0]
            if subject not in subjects:
                subjects.append(subject)
    fill_subjects_info_table(subjects, year, speciality, faculty)


def extract_from_docx(filename: str):
    """
    Parser for <filename>.docx timetable
    Process the timetable and extract speciality, studding year and subjects.
    Also refer to :code:`update_info` table to get faculty and to know if this is timetable for master's program

    :param filename: timetable filename to parse
    """
    pattern = r'([с|С|c]?.*ність)([\sа-яА-ЯіїІ"«»(),\-:;_]*?)([\s,МП]*)([\s_\-]*)([1-4]+?)([\sрн\.]+)'
    doc = Document("%s/%s" % (timetables_dir, filename))

    for table in doc.tables:
        faculty_name = get_faculty(filename)
        info = ''.join([paragraph.text for paragraph in doc.paragraphs])
        m = re.search(pattern, info)
        speciality = m.group(2)
        year = int(m.group(5))
        if is_mp(filename):
            year += 4
        for symb in ['"', "'", '«', '»', '_']:
            speciality = speciality.replace(symb, '').strip()

        subjects = list()
        for row in table.rows[1:]:
            if len(row.cells) > 2:
                subject = [run.text for run in row.cells[2].paragraphs[0].runs if run.bold and run.text]
                if subject:
                    subject = ''.join(subject)
                    subject = subject.strip()
                    if subject and subject not in subjects:
                        subjects.append(subject)
        fill_subjects_info_table(subjects, year, speciality, faculty_name)


def process_file(filename: str):
    """
    Manage to call correct parser function for giving file with filename

    :param filename: file name the will be processed
    """
    if filename.endswith('.xlsx'):
        extract_from_xlsx(filename)
    elif filename.endswith('.docx'):
        extract_from_docx(filename)
    elif filename.endswith('.doc'):
        print("Removing...    ||{:<100s}||".format(filename))
        os.remove('%s/%s' % (timetables_dir, filename))
        return

    print("Parsed...     ||{:<100s}||".format(filename))
    os.remove('%s/%s' % (timetables_dir, filename))


def extract_from_files():
    """
    Manage to process each file in /timetables directory or has just downloaded via
    :code:`timetables.timetable_downloader` function
    """
    # TODO some subjects didn't mark bold (Екологія БП-1)

    if os.listdir(timetables_dir):
        for filename in os.listdir(timetables_dir):
            process_file(filename)

    for filename in timetable_downloader.update_timetables():
        process_file(filename)


if __name__ == '__main__':
    extract_from_files()
